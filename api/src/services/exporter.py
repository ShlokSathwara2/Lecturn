import io, os, base64, re, hashlib
from typing import Optional
import httpx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from ..supabase_client import supabase


SUPABASE_URL = os.getenv("SUPABASE_URL", "")
STORAGE_BUCKET = "slide-images"


async def _fetch_image_bytes(url: str) -> Optional[bytes]:
    if not url:
        return None
    if url.startswith("/"):
        url = f"{SUPABASE_URL}/storage/v1/object/public/{STORAGE_BUCKET}{url}"
    try:
        async with httpx.AsyncClient(timeout=15) as http:
            resp = await http.get(url)
            if resp.status_code == 200:
                return resp.content
    except Exception:
        pass
    return None


def _image_ext(url: str) -> str:
    m = re.search(r"\.(jpg|jpeg|png|webp|gif|bmp)", url.lower())
    return m.group(1) if m else "png"

def _rtf_blip(ext: str) -> str:
    if ext in ("jpg", "jpeg"):
        return "jpegblip"
    return "pngblip"


# ----------------------------------------------------------------
# Data fetching
# ----------------------------------------------------------------
async def fetch_captures(subject_id: Optional[str] = None, chapter_id: Optional[str] = None) -> list[dict]:
    if chapter_id:
        chapters = [{"id": chapter_id}]
    else:
        chapters = supabase.table("chapters").select("id, title").eq("subject_id", subject_id).execute().data
    if not chapters:
        return []

    chapter_ids = [ch["id"] for ch in chapters]
    chapter_map = {ch["id"]: ch["title"] for ch in chapters}

    caps = supabase.table("captures").select(
        "id, chapter_id, raw_text, ai_content_json, image_url, cleaned_diagram_url, original_diagram_crop_url, date_taken, ai_status"
    ).in_("chapter_id", chapter_ids).order("date_taken").execute().data

    rows = []
    for cap in caps:
        ai = cap.get("ai_content_json") or {}
        enrichment = ai.get("enrichment") or {}
        rows.append({
            "id": cap["id"],
            "chapter_title": chapter_map.get(cap["chapter_id"], ""),
            "raw_text": (cap.get("raw_text") or "").strip(),
            "explanation": (enrichment.get("explanation") or "").strip(),
            "key_points": enrichment.get("key_points") or [],
            "image_url": cap.get("image_url"),
            "cleaned_diagram_url": cap.get("cleaned_diagram_url"),
            "original_diagram_crop_url": cap.get("original_diagram_crop_url"),
            "date_taken": cap.get("date_taken", ""),
            "ai_status": cap.get("ai_status", ""),
        })
    return rows


# ----------------------------------------------------------------
# TXT
# ----------------------------------------------------------------
def build_txt(rows: list[dict], title: str) -> str:
    lines = [f"{title}", "=" * len(title), ""]
    current_chapter = ""
    for r in rows:
        if r["chapter_title"] != current_chapter:
            current_chapter = r["chapter_title"]
            lines.extend([current_chapter, "-" * len(current_chapter), ""])
        if r["raw_text"]:
            lines.extend(["[Slide Content]", r["raw_text"], ""])
        if r["explanation"]:
            lines.extend(["[Study Notes]", r["explanation"], ""])
        if r["key_points"]:
            lines.append("Key Points:")
            for kp in r["key_points"]:
                lines.append(f"  - {kp}")
            lines.append("")
    return "\n".join(lines)


# ----------------------------------------------------------------
# RTF
# ----------------------------------------------------------------
def _escape_rtf(text: str) -> str:
    text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    text = text.replace("\n", "\\line ")
    return text.replace("\r", "")


def _rtf_encode_image(data: bytes, ext: str) -> str:
    BS = "\\"
    b64 = base64.b64encode(data).decode()
    blip = _rtf_blip(ext)
    return "{" + BS + "pict" + BS + blip + BS + "picwgoal4535" + BS + "pichgoal3400 " + b64 + "}"


async def build_rtf(rows: list[dict], title: str) -> bytes:
    BS = "\\"
    def esc(t):
        return _escape_rtf(t)
    def cmd(c):
        return BS + c

    header = (cmd("rtf1") + cmd("ansi") + cmd("deff0") +
              "{" + cmd("fonttbl") + "{" + cmd("f0") + " Arial;}" + "}" +
              "{" + cmd("colortbl") + ";" + cmd("red255") + cmd("green255") + cmd("blue255") + ";}" +
              "}")
    blocks = [header]
    blocks.append(cmd("pard") + cmd("b") + cmd("fs40") + cmd("qc") + " " + esc(title) + cmd("b0") + cmd("par") + cmd("par"))
    current_chapter = ""
    for r in rows:
        if r["chapter_title"] != current_chapter:
            current_chapter = r["chapter_title"]
            blocks.append(cmd("pard") + cmd("b") + cmd("fs32") + cmd("cf1") + " " + esc(current_chapter) + cmd("b0") + cmd("par") + cmd("par"))
        if r["image_url"]:
            img_data = await _fetch_image_bytes(r["image_url"])
            if img_data:
                blocks.append(cmd("pard") + cmd("qc") + " " + _rtf_encode_image(img_data, _image_ext(r["image_url"])) + cmd("par") + cmd("par"))
        if r["raw_text"]:
            blocks.append(cmd("pard") + cmd("b") + cmd("fs24") + " Slide Content" + cmd("b0") + cmd("line") + " " + esc(r["raw_text"]) + cmd("par") + cmd("par"))
        if r["explanation"]:
            blocks.append(cmd("pard") + cmd("b") + cmd("fs24") + " Study Notes" + cmd("b0") + cmd("line") + " " + esc(r["explanation"]) + cmd("par") + cmd("par"))
        if r["key_points"]:
            blocks.append(cmd("pard") + cmd("b") + cmd("fs20") + " Key Points:" + cmd("b0") + cmd("par"))
            for kp in r["key_points"]:
                blocks.append(cmd("pard") + cmd("li300") + " " + cmd("bullet") + cmd("tab") + " " + esc(kp) + cmd("par"))
            blocks.append(cmd("par"))
    blocks.append("}")
    return "".join(blocks).encode("utf-8")


# ----------------------------------------------------------------
# DOCX
# ----------------------------------------------------------------
async def build_docx(rows: list[dict], title: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, 0)
    current_chapter = ""
    for r in rows:
        if r["chapter_title"] != current_chapter:
            current_chapter = r["chapter_title"]
            doc.add_heading(current_chapter, 1)

        if r["image_url"]:
            img_data = await _fetch_image_bytes(r["image_url"])
            if img_data:
                try:
                    doc.add_picture(io.BytesIO(img_data), width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

        if r["raw_text"]:
            doc.add_heading("Slide Content", 3)
            doc.add_paragraph(r["raw_text"])

        if r["explanation"]:
            doc.add_heading("Study Notes", 3)
            doc.add_paragraph(r["explanation"])

        if r["key_points"]:
            doc.add_heading("Key Points", 3)
            for kp in r["key_points"]:
                doc.add_paragraph(kp, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ----------------------------------------------------------------
# ENEX (Evernote)
# ----------------------------------------------------------------
def _escape_xml(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


async def build_enex(rows: list[dict], title: str) -> str:
    notes_xml = []
    for r in rows:
        body_parts = []
        img_resources = []

        if r["image_url"]:
            img_data = await _fetch_image_bytes(r["image_url"])
            if img_data:
                ext = _image_ext(r["image_url"])
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif"}.get(ext, "image/png")
                b64 = base64.b64encode(img_data).decode()
                hash_hex = hashlib.md5(img_data).hexdigest()
                img_resources.append(f"""
    <resource>
      <data encoding="base64">{b64}</data>
      <mime>{mime}</mime>
      <resource-attributes>
        <file-name>image.{ext}</file-name>
      </resource-attributes>
    </resource>""")
                body_parts.append(f'<en-media type="{mime}" hash="{hash_hex}"/>')

        if r["raw_text"]:
            body_parts.append(f"<div><b>Slide Content:</b><br/>{_escape_xml(r['raw_text'])}</div>")
        if r["explanation"]:
            body_parts.append(f"<div><b>Study Notes:</b><br/>{_escape_xml(r['explanation'])}</div>")
        if r["key_points"]:
            bullets = "".join(f"<li>{_escape_xml(kp)}</li>" for kp in r["key_points"])
            body_parts.append(f"<div><b>Key Points:</b><ul>{bullets}</ul></div>")

        body = "<br/>".join(body_parts)
        created = r.get("date_taken", "")
        chapter_tag = _escape_xml(r["chapter_title"])
        note = f"""
  <note>
    <title>{_escape_xml(title)} - {chapter_tag}</title>
    <content><![CDATA[<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<!DOCTYPE en-note SYSTEM "http://xml.evernote.com/pub/enml2.dtd">
<en-note>{body}</en-note>]]></content>
    <created>{created}</created>
    <tag>{chapter_tag}</tag>{''.join(img_resources)}
  </note>"""
        notes_xml.append(note)

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE en-export SYSTEM "http://xml.evernote.com/pub/evernote-export3.dtd">
<en-export export-date="{__import__('datetime').datetime.now().isoformat()}" application="SlideScribe" version="1.0">
{''.join(notes_xml)}
</en-export>"""
    return xml


# ----------------------------------------------------------------
# PDF
# ----------------------------------------------------------------
async def build_pdf(rows: list[dict], title: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", r"C:\Windows\Fonts\arial.ttf", uni=True)
    pdf.add_font("DejaVu", "B", r"C:\Windows\Fonts\arialbd.ttf", uni=True)
    pdf.set_font("DejaVu", "B", 18)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    current_chapter = ""
    for r in rows:
        if r["chapter_title"] != current_chapter:
            current_chapter = r["chapter_title"]
            pdf.set_font("DejaVu", "B", 14)
            pdf.cell(0, 10, current_chapter, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        if r["image_url"]:
            img_data = await _fetch_image_bytes(r["image_url"])
            if img_data:
                ext = _image_ext(r["image_url"])
                fname = f"/tmp/_export_img_{r['id']}.{ext}"
                try:
                    os.makedirs("/tmp", exist_ok=True)
                    with open(fname, "wb") as f:
                        f.write(img_data)
                    pdf.image(fname, w=pdf.w - 20)
                    pdf.ln(4)
                except Exception:
                    pass
                finally:
                    try:
                        os.remove(fname)
                    except Exception:
                        pass

        if r["raw_text"]:
            pdf.set_font("DejaVu", "B", 11)
            pdf.cell(0, 7, "Slide Content", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
            pdf.multi_cell(0, 5, r["raw_text"])
            pdf.ln(2)

        if r["explanation"]:
            pdf.set_font("DejaVu", "B", 11)
            pdf.cell(0, 7, "Study Notes", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
            pdf.multi_cell(0, 5, r["explanation"])
            pdf.ln(2)

        if r["key_points"]:
            pdf.set_font("DejaVu", "B", 11)
            pdf.cell(0, 7, "Key Points:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
            for kp in r["key_points"]:
                pdf.set_x(pdf.l_margin + 5)
                pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin - 5, 5, f"- {kp}")
            pdf.ln(2)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()
